import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Historial de Proyecto: Todo Para Tu TractoCamión', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary Section
    doc.add_heading('Resumen de la Migración Técnica', level=1)
    p = doc.add_paragraph()
    p.add_run('Este documento detalla el proceso de migración de la aplicación "Todo Para Tu TractoCamión" desde una infraestructura local y Railway hacia una arquitectura moderna basada en la nube.').italic = True
    
    accomplishments = [
        "Base de Datos: Migrada exitosamente a Supabase (PostgreSQL) con 376 registros de productos.",
        "API: Desplegada en Render usando el SDK de .NET 10.0 con soporte para IPv4.",
        "Frontend: Desplegado en GitHub Pages con flujo de CI/CD mediante GitHub Actions.",
        "Correcciones: Se ajustaron esquemas de tablas, mayúsculas de columnas y carga de imágenes relativas.",
        "Seguridad: Configuración de permisos de escritura para el despliegue automático."
    ]
    for item in accomplishments:
        doc.add_paragraph(item, style='List Bullet')

    # Conversation History
    doc.add_heading('Historial de Conversaciones Recientes', level=1)
    
    conversations = [
        ("Migración de Excel a PostgreSQL", "Actualización de 376 registros con URLs de imágenes correctas y limpieza de datos."),
        ("Despliegue en Vercel y Render", "Configuración de variables de entorno y resolución de errores de compilación."),
        ("Rediseño de Interfaz", "Actualización del Hero y Header para un look más premium."),
        ("Configuración de GitHub Pages", "Activación del servicio y solución de errores 404 mediante ajuste de ramas."),
        ("Arreglo de Imágenes", "Corrección de rutas absolutas que fallaban en subdirectorios de GitHub Pages.")
    ]
    
    for title, desc in conversations:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # Image Gallery
    doc.add_heading('Galería de Imágenes del Proyecto', level=1)
    doc.add_paragraph("A continuación se presentan las capturas y evidencias enviadas durante el proceso de desarrollo:")
    
    artifact_dir = r"C:\Users\German\.gemini\antigravity\brain\1045f6da-1d74-4c21-b72b-3a5da3cb8ca4"
    image_files = [f for f in os.listdir(artifact_dir) if f.startswith("media__") and f.endswith(".png")]
    image_files.sort() # Sort by name (timestamp)
    
    for img_name in image_files:
        img_path = os.path.join(artifact_dir, img_name)
        try:
            doc.add_paragraph(f"Imagen: {img_name}")
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph("-" * 20)
        except Exception as e:
            print(f"Error al añadir imagen {img_name}: {e}")

    # Save
    output_path = r"d:\Codrava\Proyecto\Historial_TodoParaTuTractoCamion.docx"
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    create_report()
